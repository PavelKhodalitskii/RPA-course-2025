import json
from typing import Tuple
import io

import pandas as pd

from openpyxl import Workbook
from openpyxl.utils import cell as pyxl_cell
from openpyxl.chart import BarChart, Reference

from docx import Document
from docx2pdf import convert
import pdfplumber
from PyPDF2 import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from models import CarItem, CarItemsList


class CarsWorkbook:
    def __init__(self, file_path: str, header: Tuple[str], start_cell="A1"):
        self._wb = Workbook()
        self._ws = self._wb.active
        self._file_path = file_path

        self._start_cell = start_cell

        col_number, row_number = pyxl_cell.coordinate_to_tuple(start_cell)
        
        # Column pointer
        self._cp = 1
        self._start_cell_col = col_number
        # Row pointer
        self._rp = 1
        self._start_cell_row = row_number

        self._header = header
        self._paste_header()
        self._filled = False

    def _paste_header(self):
        for row in self._ws.iter_rows(min_row=self._rp, max_col=self._cp + len(self._header) - 1):
            for index, cell in enumerate(row):
                cell.value = self._header[index]
            break
        self._cp += len(self._header) - 1

    def fill_from_data(self, data: pd.DataFrame) -> None:
        self._rp += 1
        self._cp = self._start_cell_col

        self.data = data

        for index, row in enumerate(self._ws.iter_rows(min_row=self._rp, 
                                        max_row=self._rp + len(data) - 1,
                                        max_col=self._cp + len(self._header) - 1)):
            new_row = [data.index[index]] + list(data.iloc[index])

            for cell_index, cell in enumerate(row):
                cell.value = new_row[cell_index]

        self._rp += len(data) - 1
        self._cp += len(self._header)
        self._filled = True

    def fill_requiered(callable):
        def inner(self, *args, **kwargs):
            if self._filled:
                return callable(self, *args, **kwargs)
        return inner

    @fill_requiered
    def plot_hist(self, cell: str) -> None:
        chart = BarChart()
        chart.title = "Распределение цен на автомобили"
        chart.y_axis.title = 'Frequency'
        chart.x_axis.title = 'PRICE'

        data_ref = Reference(self._ws, 
                            min_col=self._start_cell_col + 2, 
                            min_row=self._start_cell_row + 1,
                            max_row=len(self.data) + 1)
        
        categories_ref = Reference(self._ws, 
                                    min_col=self._start_cell_col, 
                                    min_row=self._start_cell_row + 1, 
                                    max_row=len(self.data) + 1)
        
        self._ws.add_chart(chart, cell)
        chart.add_data(data_ref, titles_from_data=False)
        chart.set_categories(categories_ref)

    def save(self):
        self._wb.save(self._file_path)

    @property
    def file_name(self):
        return self._file_path
    
    @property
    def filled(self):
        return self._filled

class CarsDocx:
    def __init__(self, file_path):
        self._document = Document()
        self._document.add_heading('ОТЧЕТ СГЕНЕРИРОВАН РОБОТОМ!', 0)
        self._file_path = file_path

    def add_table(self, data: pd.DataFrame):
        columns = ["id", ] + list(data.columns)

        table = self._document.add_table(rows=len(data) + 1, cols=len(columns))

        for index, cell in enumerate(table.rows[0].cells):
            cell.text = columns[index]

        for index_r, row in enumerate(table.rows):
            if index_r > 0:
                for index_c, cell in enumerate(row.cells):
                    if index_c == 0:        
                        cell.text = str(data.index[index_r - 1])
                    else:
                        cell.text = str(data.iloc[index_r - 1][index_c - 1])
        
    def save(self):
        self._document.save(self._file_path)

    def to_pdf(self):
        # requires windows
        convert(self._file_path, self._file_path.replace(".docx", ".pdf"))

    @property
    def file_name(self):
        return self._file_path
    

class CarsPDFFiller:
    def __init__(self, file_path):
        self._file_path = file_path

    def set_aggr_data(self, price_sum, amount_sum):
        target_header = ["Общая сумма", "Общее кол-во"]
        fill_data = [price_sum, amount_sum]

        with pdfplumber.open(self._file_path) as pdf:
            found_table = None
            coords = None
            target_page = 0
            
            for page_idx, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                
                for table_idx, table in enumerate(tables):
                    if table:
                        header_row = table[0]
                        header_cells = [str(cell).lower() for cell in header_row if cell]
                        
                        for header in target_header:
                            if header.lower() in header_cells:
                                found_table = table
                                tables_coords = page.find_tables()
                                if table_idx < len(tables_coords):
                                    coords = tables_coords[table_idx].bbox
                                    target_page = page_idx
                                break
                        if found_table:
                            break
                if found_table:
                    break

            if found_table and coords:
                packet = io.BytesIO()
                c = canvas.Canvas(packet, pagesize=letter)
                
                c.setFont("Helvetica", 10)
                
                page_height = letter[1] 
                line_height = (coords[3] - coords[2]) / 2
                cell_width = (coords[1] - coords[0]) / 2
                
                for i, data in enumerate(fill_data):
                    x = coords[0] + cell_width * i + cell_width / 3
                    y = page_height - coords[3] + line_height / 2
                    c.drawString(x, y, str(data))
                
                c.save()
                
                packet.seek(0)
                overlay_pdf = PdfReader(packet)
                original_pdf = PdfReader(self._file_path)
                output_pdf = PdfWriter()
                
                for i, page in enumerate(original_pdf.pages):
                    if i == target_page:
                        page.merge_page(overlay_pdf.pages[0])
                    output_pdf.add_page(page)

                with open(self._file_path, "wb") as output_stream:
                    output_pdf.write(output_stream)
                    
                        
with open("data.json", "r") as data_file:
    json_data = json.load(data_file)
    cars_dict = CarItemsList.model_validate(json_data).items
    cars_df = pd.DataFrame.from_dict(cars_dict, orient="index")

    cars_frame_columns = [name for name, _ in cars_df.iloc[0]]
    cars_df.columns = cars_frame_columns
    
    for i in cars_df.columns:
        cars_df[i] = cars_df[i].apply(lambda x: x[1])

    print(cars_df.head())
    header = ["id",] + list(cars_df.columns)

    wb = CarsWorkbook(file_path="data.xlsx", header=header)
    wb.fill_from_data(cars_df)
    wb.plot_hist(cell="F1")
    wb.save()

    doc = CarsDocx(file_path="data.docx")
    doc.add_table(data=cars_df)
    doc.save()

    pdf = CarsPDFFiller(file_path="data.pdf")
    pdf.set_aggr_data(price_sum=sum(cars_df['price']), amount_sum=sum(cars_df['amount']))